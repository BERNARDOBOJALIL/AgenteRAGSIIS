"""
docx_sources.py
Carga archivos DOCX y preserva su contenido textual para ingesta RAG.
"""

from __future__ import annotations

import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator

from docx import Document as WordDocument
from docx.document import Document as WordProcessingDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document

log = logging.getLogger(__name__)

DEFAULT_DOCX_PATTERN = "*.docx"
_EXCLUDED_PATH_PARTS = {".venv", "chroma_db", "__pycache__"}


def _normalize_line(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").replace("\xa0", " ")).strip()


def _heading_prefix(style_name: str) -> str:
    match = re.search(r"(\d+)", style_name or "")
    if not match:
        return "#"
    level = max(1, min(6, int(match.group(1))))
    return "#" * level


def _iter_block_items(document: WordProcessingDocument) -> Iterator[Paragraph | Table]:
    """Itera parrafos y tablas en el orden real del documento."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _extract_table_lines(table: Table, table_idx: int) -> list[str]:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_normalize_line(cell.text).replace("|", "/") for cell in row.cells]
        if any(cells):
            rows.append(cells)

    if not rows:
        return []

    width = max(len(row) for row in rows)
    padded_rows = [row + [""] * (width - len(row)) for row in rows]

    lines = [f"[TABLA {table_idx}]", "| " + " | ".join(cell or " " for cell in padded_rows[0]) + " |"]
    lines.append("| " + " | ".join("---" for _ in range(width)) + " |")
    for row in padded_rows[1:]:
        lines.append("| " + " | ".join(cell or " " for cell in row) + " |")
    return lines


def _datetime_to_iso(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.astimezone(timezone.utc).isoformat()
    return _normalize_line(str(value))


def _extract_docx_text(word: WordProcessingDocument) -> tuple[str, dict[str, int]]:

    lines: list[str] = []
    paragraph_count = 0
    table_count = 0
    table_rows = 0

    for block in _iter_block_items(word):
        if isinstance(block, Paragraph):
            text = _normalize_line(block.text)
            if not text:
                continue
            paragraph_count += 1

            style_name = getattr(getattr(block, "style", None), "name", "") or ""
            if style_name.lower().startswith("heading"):
                lines.append(f"{_heading_prefix(style_name)} {text}")
            else:
                lines.append(text)
            continue

        table_count += 1
        table_rows += len(block.rows)
        lines.extend(_extract_table_lines(block, table_count))

    for section in word.sections:
        for paragraph in section.header.paragraphs:
            text = _normalize_line(paragraph.text)
            if text:
                lines.append(f"[ENCABEZADO] {text}")
        for paragraph in section.footer.paragraphs:
            text = _normalize_line(paragraph.text)
            if text:
                lines.append(f"[PIE_PAGINA] {text}")

    return "\n".join(lines).strip(), {
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "table_rows": table_rows,
    }


def load_docx_documents(docx_dir: str, pattern: str = DEFAULT_DOCX_PATTERN) -> list[Document]:
    """Carga todos los DOCX encontrados y devuelve documentos listos para normalizar."""
    base_path = Path(docx_dir)
    if not base_path.exists():
        log.warning("La carpeta DOCX '%s' no existe.", docx_dir)
        return []

    files = []
    for candidate in sorted(base_path.rglob(pattern), key=lambda p: str(p).lower()):
        if any(part in _EXCLUDED_PATH_PARTS for part in candidate.parts):
            continue
        if candidate.is_file():
            files.append(candidate)

    if not files:
        log.warning("No se encontraron DOCX en '%s' con patron '%s'.", docx_dir, pattern)
        return []

    documents: list[Document] = []
    for file_path in files:
        try:
            word = WordDocument(str(file_path))
            content, stats = _extract_docx_text(word)
        except Exception as exc:
            log.warning("Error leyendo DOCX '%s': %s", file_path.name, exc)
            continue

        if not content:
            log.warning("DOCX sin contenido util: %s", file_path.name)
            continue

        core = word.core_properties
        file_stat = file_path.stat()
        source_last_modified_utc = datetime.fromtimestamp(file_stat.st_mtime, tz=timezone.utc).isoformat()
        metadata = {
            "fuente": "docx_local",
            "categoria": "docx",
            "title": file_path.stem,
            "source_file": str(file_path.resolve()),
            "source_filename": file_path.name,
            "source_ext": ".docx",
            "source_last_modified_utc": source_last_modified_utc,
            "docx_size_bytes": file_stat.st_size,
            "docx_paragraphs": stats["paragraph_count"],
            "docx_tables": stats["table_count"],
            "docx_table_rows": stats["table_rows"],
            "docx_author": _normalize_line(core.author or ""),
            "docx_subject": _normalize_line(core.subject or ""),
            "docx_category": _normalize_line(core.category or ""),
            "docx_keywords": _normalize_line(core.keywords or ""),
            "docx_created_utc": _datetime_to_iso(core.created),
            "docx_modified_utc": _datetime_to_iso(core.modified),
            "fecha_normalizacion_utc": source_last_modified_utc,
        }
        metadata = {k: v for k, v in metadata.items() if v not in (None, "")}
        documents.append(Document(page_content=content, metadata=metadata))

    log.info("DOCX convertidos a texto: %s archivo(s).", len(documents))
    return documents
